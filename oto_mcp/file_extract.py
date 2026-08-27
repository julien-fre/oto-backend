"""Extraire le TEXTE d'un fichier déposé — pour qu'on le trouve par ce qu'il contient (#298).

Un fichier déposé dans un projet n'était trouvable que par son nom, son titre ou sa
description. Un PDF de trente pages était donc, du point de vue de la recherche, un
nom de fichier : mal nommé, introuvable — alors que tout le reste du contenu (pages,
tableaux, guides, procédures) est cherchable plein texte. Le fichier était le seul
angle mort, et c'est souvent lui qui porte la matière apportée par le client.

## Ce que ce module fait, et ce qu'il ne fait pas

**Il rend du texte, ou il dit pourquoi il n'en rend pas.** C'est tout. Il ne parle ni
à la base, ni à S3, ni au réseau : on lui donne des octets, il rend un `Extraction`.
Cette frontière est ce qui le rend testable sans rien monter, et ce qui permet de
l'appeler depuis un worker hors de la boucle d'événements (le serveur est mono-loop —
une extraction synchrone dans la boucle la gèlerait, cf. `docs/event-loop-perf.md`).

**Il n'y a PAS d'OCR.** Lire une image est un autre métier et un autre coût ; un fichier
image ressort `unsupported`, ce qui est une réponse honnête. Le jour où l'OCR se
décide, il s'ajoute ici sans rien changer au reste.

## Les statuts sont TERMINAUX, et c'est le point de conception

Un fichier chiffré n'est pas un échec transitoire : le retenter à chaque passage du
worker coûte pour toujours et ne réussira jamais. Chaque cas se nomme donc
(`unsupported`, `encrypted`, `empty`, `failed`), et seul `failed` — l'imprévu — laisse
la porte ouverte à une reprise. Un statut nommé est aussi ce qui rend l'interface
honnête : « format non supporté » n'est pas « en cours ».

## Coût mesuré (banc du 13/08, 6 documents réels : decks, plaquette, supports)

**34 pages, 44 050 caractères, 0,78 s** — soit ~130 ms par document. Un document fait
~7 300 caractères, c'est-à-dire **quelques pages rédigées, pas cent** : le pronostic
« deux ordres de grandeur au-dessus d'une page » était faux d'un facteur dix. Indexer
ce volume coûte 3 ms par fichier déposé et 0,5 Mo pour 500 fichiers.

⚠️ Cet échantillon penche vers le PEU de texte (des slides portent ~1 200 caractères
par page, un rapport en prose 3 à 5 fois plus) : les chiffres sont un **plancher**.
Même multipliés par cinq, ils restent modestes — c'est ce qui a écarté l'idée d'un
process d'extraction séparé.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from typing import Optional

log = logging.getLogger(__name__)

# Bornes du texte retenu. La première protège la base d'un fichier pathologique (un
# export de log de 500 Mo n'a pas à devenir une ligne) ; la seconde évite d'écrire un
# statut « extrait » pour trois caractères de garde d'un PDF scanné, qui est en
# réalité un cas OCR déguisé.
MAX_TEXT_CHARS = 2_000_000
MIN_TEXT_CHARS = 16

# ⚠️ **Borne de DÉCOMPRESSION — c'est une protection, pas un réglage.** Un `.docx`
# comme un `.xlsx` est une archive ZIP : la taille du fichier reçu ne dit RIEN de ce
# qu'il pèse une fois ouvert. Mesuré le 13/08 sur ce module : une archive de **400 ko**
# faisait monter le process à **638 Mo** (facteur 1 600) — le document entier est
# parsé en mémoire AVANT que `MAX_TEXT_CHARS` ne tronque quoi que ce soit, donc la
# borne de sortie n'protège de rien.
#
# Le serveur est MONO-LOOP : un dépassement mémoire ne dégrade pas une requête, il tue
# le process et toutes les sessions avec. Et le fichier vient d'un upload utilisateur,
# donc l'entrée est hostile par construction. On refuse donc AVANT d'ouvrir, en lisant
# la taille annoncée par le catalogue du ZIP.
MAX_UNCOMPRESSED_BYTES = 64 * 1024 * 1024

# Statuts. `ok` mis à part, tous sont TERMINAUX sauf `failed` (cf. `is_retryable`).
OK = "ok"
UNSUPPORTED = "unsupported"      # format qu'on ne sait pas lire (image, archive, binaire)
ENCRYPTED = "encrypted"          # protégé par mot de passe — ne réussira jamais seul
EMPTY = "empty"                  # lu, mais sans texte utile (scan sans OCR, doc vide)
TOO_LARGE = "too_large"          # décompresserait au-delà de la borne — refusé AVANT lecture
REJECTED_DTD = "rejected_dtd"    # XML à déclaration d'entités — refusé AVANT parsing
FAILED = "failed"                # imprévu (fichier tronqué, lib qui lève) — reprenable


@dataclass(frozen=True)
class Extraction:
    """Le résultat, qu'il ait abouti ou non.

    `status` vaut toujours quelque chose ; `text` n'est rempli que sur `ok`. `detail`
    porte de quoi comprendre sans rouvrir le fichier (le type d'erreur, jamais son
    contenu — un message de lib peut recracher des octets du document)."""
    status: str
    text: str = ""
    pages: Optional[int] = None
    detail: str = ""

    @property
    def ok(self) -> bool:
        return self.status == OK


def is_retryable(status: str) -> bool:
    """Seul l'imprévu se retente. Un format non supporté ou un fichier chiffré ne
    changera pas d'avis au prochain passage du worker — les retenter, c'est payer une
    file qui ne se vide jamais."""
    return status == FAILED


# Extensions → extracteur. On route sur l'EXTENSION et non sur le mime déclaré : le
# mime d'un upload vient du client, qui se trompe couramment (`application/octet-stream`
# sur un PDF parfaitement valide). Le mime sert de repli quand l'extension manque.
_BY_EXT = {}


def _register(*exts):
    def deco(fn):
        for e in exts:
            _BY_EXT[e] = fn
        return fn
    return deco


@_register("txt", "md", "markdown", "csv", "json", "log", "rst")
def _extract_text(data: bytes) -> Extraction:
    """Texte brut. `errors='replace'` plutôt que de lever : un octet douteux au milieu
    d'un fichier de 10 000 lignes ne doit pas rendre les 9 999 autres introuvables."""
    return Extraction(OK, data.decode("utf-8", errors="replace"))


def _join_bounded(parts) -> tuple:
    """Concaténer en s'ARRÊTANT à la borne, au lieu d'accumuler puis tronquer.

    La nuance n'est pas cosmétique : un document à dix mille pages tient tout entier
    en mémoire avant d'être coupé si l'on accumule d'abord. On coupe donc pendant.
    Rend `(texte, tronqué)`."""
    total, morceaux = 0, []
    for p in parts:
        if not p:
            continue
        morceaux.append(p)
        total += len(p)
        if total > MAX_TEXT_CHARS:
            return "\n".join(morceaux)[:MAX_TEXT_CHARS], True
    return "\n".join(morceaux), False


@_register("pdf")
def _extract_pdf(data: bytes) -> Extraction:
    from pypdf import PdfReader
    from pypdf.errors import FileNotDecryptedError, PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
        if getattr(reader, "is_encrypted", False):
            # Un PDF « chiffré » sans mot de passe s'ouvre avec une chaîne vide — le
            # cas est courant (protection en écriture seule). On ne renonce donc
            # qu'après avoir essayé.
            try:
                reader.decrypt("")
            # noqa: SILENT — PDF protégé : le verdict ENCRYPTED est rendu à l'appelant
            except Exception:
                return Extraction(ENCRYPTED, detail="pdf protégé par mot de passe")
        n_pages = len(reader.pages)
        # Générateur : les pages sont extraites AU FIL de la concaténation, qui
        # s'arrête à la borne — un PDF à dix mille pages ne tient jamais en entier.
        texte, tronque = _join_bounded(
            (p.extract_text() or "") for p in reader.pages)
    except FileNotDecryptedError:
        return Extraction(ENCRYPTED, detail="pdf protégé par mot de passe")
    # noqa: SILENT — verdict FAILED + type d'exception rendus à l'appelant
    except (PdfReadError, Exception) as e:            # tronqué, corrompu, inattendu
        return Extraction(FAILED, detail=type(e).__name__)
    return Extraction(OK, texte, pages=n_pages,
                      detail=f"tronqué à {MAX_TEXT_CHARS} caractères" if tronque else "")


# Marqueurs d'une déclaration d'entités XML. Un document bureautique LÉGITIME n'en
# porte jamais : les producteurs (Word, Excel, LibreOffice, les libs de génération)
# écrivent du XML sans DTD. Leur présence n'est donc pas un cas limite à gérer, c'est
# un document forgé.
_DTD_MARQUEURS = (b"<!DOCTYPE", b"<!ENTITY", b"<!doctype", b"<!entity")


def _zip_declares_entities(data: bytes) -> Optional[Extraction]:
    """Refuser un document dont une partie XML déclare des entités — AVANT de parser.

    ## Pourquoi cette garde existe alors que la mesure dit qu'on est déjà protégé

    Vérifié le 13/08 en fabriquant les attaques : ni `python-docx`/lxml ni `openpyxl`
    ne résolvent l'entité externe (aucune fuite de fichier local, aucune requête
    sortante) et l'expansion d'entités est refusée (mémoire stable). Le risque signalé
    n'était pas exploitable.

    Mais cette protection est **empruntée** : elle appartient au parseur des
    bibliothèques, pas à nous. Elle tient tant qu'elles ne changent pas de moteur XML,
    ce dont personne ne nous préviendra — et le jour où ça arrive, la faille revient
    par une mise à jour de routine, sur un chemin qui traite des fichiers hostiles.
    Le pré-scan, lui, ne dépend d'aucune bibliothèque : il ferme la classe entière,
    quel que soit ce qui parse derrière.

    C'est aussi pourquoi il est préféré à `defusedxml` : la dépendance ne corrige rien
    de mesurable ici (les deux moteurs résistent déjà) et ne couvrirait que le lecteur
    qui la consulte, alors que cette garde couvre tout format ZIP+XML — celui qu'on
    ajoutera demain compris.

    ⚠️ Le contrôle vient APRÈS `_zip_too_large` : lire les parties suppose de les
    décompresser, donc la borne de taille doit avoir déjà parlé.
    """
    import zipfile
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            for it in z.infolist():
                if not it.filename.lower().endswith((".xml", ".rels")):
                    continue
                contenu = z.read(it.filename)
                if any(m in contenu for m in _DTD_MARQUEURS):
                    return Extraction(
                        REJECTED_DTD,
                        detail=f"déclaration d'entités XML dans « {it.filename} »")
    # noqa: SILENT — pas un zip lisible : la lib rendra `failed`
    except Exception:
        return None                       # pas un zip lisible : la lib dira `failed`
    return None


def _zip_too_large(data: bytes) -> Optional[Extraction]:
    """Refuser AVANT d'ouvrir une archive qui décompresserait au-delà de la borne.

    On lit le CATALOGUE du ZIP (`infolist`), qui annonce la taille décompressée de
    chaque membre : ça ne décompresse rien, donc le contrôle ne peut pas être la
    victime de ce qu'il contrôle.

    ⚠️ Une taille annoncée peut MENTIR (un en-tête forgé). C'est acceptable ici : le
    mensonge ne peut que sous-déclarer, et la lib s'arrête alors sur une archive
    incohérente (`failed`). Ce qu'on ferme, c'est le cas simple et efficace — une
    archive honnête et énorme —, celui qui coûte 400 ko à l'attaquant et 638 Mo au
    serveur. Une borne dure sur la mémoire du process est un autre sujet
    (`docs/event-loop-perf.md`), pas celui de ce module.
    """
    import zipfile
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            total = sum(i.file_size for i in z.infolist())
    # noqa: SILENT — pas un zip lisible : la lib rendra `failed`
    except Exception:
        return None                       # pas un zip lisible : la lib dira `failed`
    if total > MAX_UNCOMPRESSED_BYTES:
        return Extraction(
            TOO_LARGE,
            detail=f"décompresserait {total // 1048576} Mo "
                   f"(borne {MAX_UNCOMPRESSED_BYTES // 1048576} Mo)")
    return None


def _zip_guards(data: bytes) -> Optional[Extraction]:
    """Les gardes d'une archive ZIP+XML, dans l'ORDRE qui les rend sûres.

    Une seule porte pour tous les formats de cette famille : un lecteur ajouté demain
    (`.pptx`, `.odt`) l'appelle et hérite des deux protections, au lieu d'en oublier
    une. L'ordre n'est pas indifférent — le scan des entités lit les parties, donc il
    suppose que la borne de taille a déjà refusé les archives démesurées."""
    return _zip_too_large(data) or _zip_declares_entities(data)


@_register("docx")
def _extract_docx(data: bytes) -> Extraction:
    from docx import Document

    refus = _zip_guards(data)
    if refus is not None:
        return refus
    try:
        doc = Document(io.BytesIO(data))
        # Les paragraphes ET les tableaux : dans un document bureautique, une part
        # notable du contenu utile (chiffres, engagements) vit dans les tableaux, que
        # `doc.paragraphs` seul ne voit pas.
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                parts.extend(c.text for c in row.cells)
        texte, tronque = _join_bounded(parts)
    # noqa: SILENT — verdict FAILED + type d'exception rendus à l'appelant
    except Exception as e:
        return Extraction(FAILED, detail=type(e).__name__)
    return Extraction(OK, texte,
                      detail=f"tronqué à {MAX_TEXT_CHARS} caractères" if tronque else "")


@_register("xlsx")
def _extract_xlsx(data: bytes) -> Extraction:
    from openpyxl import load_workbook

    refus = _zip_guards(data)
    if refus is not None:
        return refus
    try:
        # `read_only` + `data_only` : on veut les VALEURS, pas les formules, et sans
        # charger le classeur entier en mémoire (un export de 50 000 lignes existe).
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)

        def _cellules():
            for ws in wb.worksheets:
                yield str(ws.title)
                for row in ws.iter_rows(values_only=True):
                    for v in row:
                        if v is not None:
                            yield str(v)

        texte, tronque = _join_bounded(_cellules())
        wb.close()
    # noqa: SILENT — verdict FAILED + type d'exception rendus à l'appelant
    except Exception as e:
        return Extraction(FAILED, detail=type(e).__name__)
    return Extraction(OK, texte,
                      detail=f"tronqué à {MAX_TEXT_CHARS} caractères" if tronque else "")


def supported_extensions() -> set:
    """Ce qu'on sait lire — dérivé du registre, jamais recopié : une liste en double
    finit par mentir (celle de la doc, celle du code, et la vraie)."""
    return set(_BY_EXT)


_MIME_EXT = {"application/pdf": "pdf", "text/plain": "txt", "text/markdown": "md",
             "text/csv": "csv", "application/json": "json"}


def _extension_of(filename: str, mime: str = "") -> str:
    """L'extension qui décide du lecteur — le mime n'est qu'un REPLI, et seulement
    quand l'extension MANQUE.

    ⚠️ La nuance a été trouvée en écrivant les tests du worker, et elle compte : le
    repli ne doit pas s'appliquer à une extension présente mais non supportée. Sinon
    un fichier nommé `.png`, déclaré `application/pdf` par un client qui se trompe (ou
    qui ment), part chez le lecteur PDF — alors que le principe posé est que
    l'extension fait autorité. Le mime vient du client ; il ne peut pas servir à
    contredire ce que le fichier dit de lui-même, seulement à combler un silence.
    """
    nom = filename or ""
    ext = nom.rsplit(".", 1)[-1].lower() if "." in nom else ""
    if ext:
        return ext                      # présente : elle décide, supportée ou non
    # Sans extension (un upload d'API, un `blob`), le mime est tout ce qu'on a.
    return _MIME_EXT.get((mime or "").split(";")[0].strip().lower(), "")


def extract(data: bytes, filename: str, mime: str = "") -> Extraction:
    """Le texte d'un fichier, ou la raison nommée de son absence.

    Ne lève JAMAIS : ce module est appelé par un worker de fond qui traite une file.
    Une exception qui remonte, c'est un fichier qui bloque la file ou un worker qui
    meurt — alors qu'un `failed` nommé laisse la file avancer et le cas visible.
    """
    if not data:
        return Extraction(EMPTY, detail="fichier vide")
    ext = _extension_of(filename, mime)
    fn = _BY_EXT.get(ext)
    if fn is None:
        return Extraction(UNSUPPORTED, detail=f"format « {ext or 'inconnu'} » non supporté")
    try:
        out = fn(data)
    except Exception as e:                    # ceinture : un extracteur ne doit pas tuer la file
        log.warning("extraction %s a levé : %s", ext, type(e).__name__)
        return Extraction(FAILED, detail=type(e).__name__)
    if not out.ok:
        return out

    text = out.text.strip()
    if len(text) < MIN_TEXT_CHARS:
        # Le cas le plus courant ici est le PDF SCANNÉ : la lecture réussit et ne rend
        # rien, parce que le texte est une image. Le nommer `empty` (et non `ok` avec
        # un texte vide) est ce qui permet à l'interface de le dire, et à un futur
        # lot d'OCR de retrouver exactement la population concernée.
        return Extraction(EMPTY, pages=out.pages,
                          detail="aucun texte extractible (document scanné ?)")
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS]
        return Extraction(OK, text, pages=out.pages,
                          detail=f"tronqué à {MAX_TEXT_CHARS} caractères")
    return Extraction(OK, text, pages=out.pages, detail=out.detail)
